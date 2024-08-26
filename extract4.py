import pypandoc
# from google.colab import files
import docx
from docx.shared import Pt
import requests
# Example translation function (replace this with your actual translation logic)
def translate_text(text):
    # Simulated translation (replace this with actual translation logic)
    return f"{text}"

# Step 1: Convert the document using pypandoc
input_file = r"C:\Users\Shivani Priya R\Downloads\gopi-main\gopi-main\CISO Function Policies\1. ISMS Policies\Asset Management Policy 1.0.docx"
output_file = 'recreated_document.docx'
output = pypandoc.convert_file(input_file, 'docx', outputfile=output_file)
assert output == ""  # Ensure the conversion was successful

# Step 2: Open the converted document with python-docx
doc = docx.Document(input_file)

# Step 3: Translate and replace each run in paragraphs, tables, headers, and footers
# def translate_text(text_chunk: str) -> str:

#     # Map the frontend language to the corresponding FLORES-200 code
#     # target_lang = language_mapping.get(frontend_lang.lower())
    
#     # if not target_lang:
#     #     raise ValueError(f"Language '{frontend_lang}' is not supported or mapping is missing.")
    
#     # Define the endpoint URL (replace with your actual Ngrok URL)
#     url = "https://rich-concrete-perch.ngrok-free.app/translate"
    
#     # Prepare the request payload
#     payload = {
#         "text": text_chunk,
#         "target_lang": "fra_Latn"
#     }
    
#     # Send the POST request to the FastAPI endpoint
#     response = requests.post(url, json=payload)
    
#     # Check if the request was successful
#     if response.status_code == 200:
#         translation = response.json().get('translation')
#         return translation
#     else:
#         raise Exception(f"Failed to translate text. Status code: {response.status_code}, Detail: {response.text}")
def translate_text(text_chunk: str) -> str:

    # Map the frontend language to the corresponding FLORES-200 code
    # target_lang = language_mapping.get(frontend_lang.lower())
    
    # if not target_lang:
    #     raise ValueError(f"Language '{frontend_lang}' is not supported or mapping is missing.")
    
    # Define the endpoint URL (replace with your actual Ngrok URL)
    url = "https://rich-concrete-perch.ngrok-free.app/translate"
    
    # Prepare the request payload
    payload = {
        "text": text_chunk,
        "target_lang": "fra_Latn"
    }
    
    # Send the POST request to the FastAPI endpoint
    response = requests.post(url, json=payload)
    
    # Check if the request was successful
    if response.status_code == 200:
        translation = response.json().get('translation')
        return translation
    else:
        raise Exception(f"Failed to translate text. Status code: {response.status_code}, Detail: {response.text}")



# Function to translate paragraphs (handles both normal paragraphs and those in tables)
def translate_paragraphs(paragraphs):
    for para in paragraphs:
        # Perform translation
        original_text = para.text
        # print(original_text)
        trans = translate_text(original_text)
        print(original_text+"    "+trans)
        for run in para.runs:
            # Capture original style
            font_name = run.font.name
            font_size = run.font.size
            bold = run.bold
            italic = run.italic
            underline = run.underline
            color = run.font.color.rgb


             # Retain the original style while replacing the text

            # Restore original style
            run.font.name = font_name
            run.font.size = font_size
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if color:
                run.font.color.rgb = color

            para.text=trans
            # run = translated_text
            # break



# Handling main body paragraphs
translate_paragraphs(doc.paragraphs)

# Handling tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            translate_paragraphs(cell.paragraphs)

# Handling headers
for section in doc.sections:
    header = section.header
    translate_paragraphs(header.paragraphs)

# Handling footers
for section in doc.sections:
    footer = section.footer
    translate_paragraphs(footer.paragraphs)

# Handling images (copying images as is)
for shape in doc.inline_shapes:
    shape._inline.graphic.graphicData = shape._inline.graphic.graphicData

# Handling borders (copying border settings as is)
for table in doc.tables:
    tbl = table._element
    tblBorders = tbl.xpath(".//w:tblBorders")
    for border in tblBorders:
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            sideBorder = border.find(f'w:{side}', tbl.nsmap)
            if sideBorder is not None:
                # Modify border properties here if needed
                pass

# Step 4: Save the updated document
updated_file = 'updated_newdocument.docx'
doc.save(updated_file)

# Step 5: Download the updated document
# files.download(updated_file)
